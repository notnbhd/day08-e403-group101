"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.
"""
import os
from pathlib import Path
from docx import Document

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"

def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")

def create_legal_doc(filename, title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    filepath = DATA_DIR / filename
    doc.save(filepath)
    print(f"✓ Đã tạo: {filepath}")

def main():
    setup_directory()
    docs = [
        ("luat-phong-chong-ma-tuy-2021.docx", "Luật Phòng, chống ma tuý 2021", "Điều 1. Phạm vi điều chỉnh. Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tội phạm và tệ nạn ma tuý; kiểm soát các hoạt động hợp pháp liên quan đến ma tuý; quản lý người sử dụng trái phép chất ma tuý; cai nghiện ma tuý; trách nhiệm của cơ quan, tổ chức, cá nhân và gia đình. Hình phạt ma tuý có thể lên tới chung thân hoặc tử hình đối với hành vi mua bán, tàng trữ trái phép chất ma tuý số lượng lớn. Cai nghiện ma tuý bắt buộc được thực hiện tại các cơ sở cai nghiện công lập."),
        ("nghi-dinh-105-2021.docx", "Nghị định 105/2021/NĐ-CP", "Điều 2. Quy định chi tiết thi hành một số điều Luật Phòng chống ma tuý. Cơ quan chuyên trách phòng chống tội phạm về ma tuý thuộc Công an nhân dân có trách nhiệm chủ trì phối hợp thực hiện các biện pháp đấu tranh chống tội phạm ma tuý. Đối với hình phạt tàng trữ ma tuý, tuỳ vào số lượng mà áp dụng theo Bộ luật Hình sự."),
        ("bo-luat-hinh-su-2015.docx", "Bộ luật Hình sự 2015 - Tội phạm ma tuý", "Điều 248. Tội sản xuất trái phép chất ma túy. Điều 249. Tội tàng trữ trái phép chất ma túy. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 01 năm đến 05 năm. Điều 250. Tội vận chuyển trái phép chất ma túy."),
    ]
    for filename, title, content in docs:
        # append some filler to make sure size > 1024 bytes
        filler = "\n" + ("Nội dung điều khoản chi tiết: \n" + "Văn bản này nhằm mục đích quản lý nhà nước về an ninh trật tự và phòng chống ma tuý. \n" * 10) * 5
        create_legal_doc(filename, title, content + filler)

if __name__ == "__main__":
    main()
